from __future__ import annotations

from pathlib import Path
import re
import shutil

from PIL import Image, ImageDraw, ImageFont, JpegImagePlugin
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

import build_v12_english_sci_package as base
from compose_figure1_npj_multipanel import compose as compose_figure1_npj


ROOT = Path(__file__).resolve().parents[1]
NPJ_DST = ROOT / "13_manuscript_and_submission" / "NPJ_Biofilms_submission_package_20260524"
MAIN_DOC = "Manuscript_Dietary_inflammatory_potential_oral_microbiome_npj.docx"
SUPP_DOC = "Supplementary_Information_Diet_oral_microbiome_npj.docx"
COVER_DOC = "Cover_letter_npj_Biofilms_and_Microbiomes.docx"
SUPP_DATA = "Supplementary_Data_1_numeric_tables.xlsx"
SOURCE_DATA = "Source_Data.xlsx"

base.DST = NPJ_DST
base.FINAL_DST = NPJ_DST
base.MAIN_DOC = MAIN_DOC
base.SUPP_DOC = SUPP_DOC


TITLE = "Dietary inflammatory potential and oral microbiome ecology in NHANES"
ABSTRACT = (
    "Dietary exposures may shape oral biofilms, but population-scale evidence linking inflammatory diet patterns "
    "to oral microbial ecology remains limited. We analyzed NHANES 2009-2012 Oral Microbiome Project data to test "
    "associations of the energy-adjusted Dietary Inflammatory Index, the unadjusted Dietary Inflammatory Index, "
    "and six diet-quality indices with oral microbiome features. Among 2799 complete-case participants for full "
    "models, alpha-diversity associations were weak and did not survive false-discovery-rate correction. "
    "Beta-diversity associations were small but recurrent across diet indices and distance metrics, with PERMANOVA "
    "R2 values from 0.00031 to 0.00158 and no tests meeting FDR < 0.05. Taxon-level models and co-abundance modules "
    "prioritized candidate taxa and ecological modules, including a module signal associated with the healthful plant-based diet index. These findings "
    "indicate small but consistent ecological signals rather than large compositional shifts, and they require "
    "contemporary longitudinal and functional validation."
)


FIGURE_LEGENDS = [
    (
        "Figure 1",
        "Analytic sample construction and diet-index architecture. a) Sample construction from oral microbiome "
        "metadata to the complete-case full-model analytic sample. b) Correlation network across the eight diet "
        "indices. c) Principal-component loading map showing the shared diet-quality axis and the secondary "
        "inflammatory or plant-based axis. d) Analysis hierarchy linking primary exposure, microbiome layers and "
        "interpretation tiers. e) Survey-weighted associations of diet principal components with alpha-diversity "
        "outcomes. f) Exploratory integrated evidence map across analytical layers. E-DII, energy-adjusted Dietary "
        "Inflammatory Index; FDR, false-discovery rate."
    ),
    (
        "Figure 2",
        "Microbiome diversity and dispersion analyses. The figure summarizes alpha-diversity associations, "
        "nonlinear diagnostics, PERMANOVA R2 values, PCoA visualizations, constrained ordination summaries and "
        "PERMDISP diagnostics. Beta-diversity signals were small and should be interpreted with the PERMDISP "
        "diagnostics because some community-structure findings may include dispersion components."
    ),
    (
        "Figure 3",
        "Taxon-level internal method-concordance and candidate taxa. a) Matrix of survey-weighted centered "
        "log-ratio taxonomic-feature associations across diet indices. b) Ranked lollipop plot of candidate taxa and "
        "dose-response support. These panels are used for exploratory prioritization and do not constitute external "
        "validation."
    ),
    (
        "Figure 4",
        "Exploratory co-abundance modules and integrated ecological evidence. a) Strong co-abundance network edges "
        "among selected genera and modules. b) Split-tile evidence map summarizing repeated signals across taxa, "
        "modules and sensitivity analyses. Co-abundance modules should be interpreted as exploratory ecological "
        "summaries rather than direct microbial interactions."
    ),
]


NATURE_REFERENCES = [
    "Dewhirst, F. E. et al. The human oral microbiome. J. Bacteriol. 192, 5002-5017 (2010).",
    "Lamont, R. J., Koo, H. & Hajishengallis, G. The oral microbiota: dynamic communities and host interactions. Nat. Rev. Microbiol. 16, 745-759 (2018).",
    "Baker, J. L., Mark Welch, J. L., Kauffman, K. M., McLean, J. S. & He, X. The oral microbiome: diversity, biogeography and human health. Nat. Rev. Microbiol. 22, 89-104 (2024).",
    "Shivappa, N., Steck, S. E., Hurley, T. G., Hussey, J. R. & Hebert, J. R. Designing and developing a literature-derived, population-based dietary inflammatory index. Public Health Nutr. 17, 1689-1696 (2014).",
    "Krebs-Smith, S. M. et al. Update of the Healthy Eating Index: HEI-2015. J. Acad. Nutr. Diet. 118, 1591-1602 (2018).",
    "Fung, T. T. et al. Adherence to a DASH-style diet and risk of coronary heart disease and stroke in women. Arch. Intern. Med. 168, 713-720 (2008).",
    "Fung, T. T. et al. Diet-quality scores and plasma concentrations of markers of inflammation and endothelial dysfunction. Am. J. Clin. Nutr. 82, 163-173 (2005).",
    "Satija, A. et al. Healthful and unhealthful plant-based diets and the risk of coronary heart disease in U.S. adults. J. Am. Coll. Cardiol. 70, 411-422 (2017).",
    "Willett, W. et al. Food in the Anthropocene: the EAT-Lancet Commission on healthy diets from sustainable food systems. Lancet 393, 447-492 (2019).",
    "Cacau, L. T., De Carli, E. & de Carvalho, A. M. Development and validation of an index based on EAT-Lancet recommendations: the Planetary Health Diet Index. Nutrients 13, 1698 (2021).",
    "National Center for Health Statistics. NHANES 2009-2012 Oral Microbiome Project: data documentation, data files and variable/annotation files. Centers for Disease Control and Prevention. https://wwwn.cdc.gov/nchs/nhanes/omp/Default.aspx (accessed 25 May 2026).",
    "Human Microbiome Project Consortium. Structure, function and diversity of the healthy human microbiome. Nature 486, 207-214 (2012).",
    "National Center for Health Statistics. National Health and Nutrition Examination Survey: analytic guidelines. Centers for Disease Control and Prevention. https://wwwn.cdc.gov/nchs/nhanes/analyticguidelines.aspx (accessed 25 May 2026).",
    "Lozupone, C. & Knight, R. UniFrac: a new phylogenetic method for comparing microbial communities. Appl. Environ. Microbiol. 71, 8228-8235 (2005).",
    "Anderson, M. J. A new method for non-parametric multivariate analysis of variance. Austral Ecol. 26, 32-46 (2001).",
    "Anderson, M. J. Distance-based tests for homogeneity of multivariate dispersions. Biometrics 62, 245-253 (2006).",
    "Oksanen, J. et al. vegan: Community Ecology Package. R package documentation. https://cran.r-project.org/package=vegan (accessed 25 May 2026).",
    "Aitchison, J. The Statistical Analysis of Compositional Data (Chapman and Hall, 1986).",
    "Segata, N. et al. Metagenomic biomarker discovery and explanation. Genome Biol. 12, R60 (2011).",
    "Lin, H. & Peddada, S. D. Analysis of compositions of microbiomes with bias correction. Nat. Commun. 11, 3514 (2020).",
    "Lin, H. & Das Peddada, S. Multigroup analysis of compositions of microbiomes with covariate adjustments and repeated measures. Nat. Methods 20, 1655-1664 (2023).",
    "Mallick, H. et al. Multivariable association discovery in population-scale meta-omics studies. PLoS Comput. Biol. 17, e1009442 (2021).",
    "Benjamini, Y. & Hochberg, Y. Controlling the false discovery rate: a practical and powerful approach to multiple testing. J. R. Stat. Soc. Series B 57, 289-300 (1995).",
    "Harrell, F. E. Jr. Regression Modeling Strategies 2nd edn (Springer, 2015).",
    "Lumley, T. Complex Surveys: A Guide to Analysis Using R (Wiley, 2010).",
    "VanderWeele, T. J. Explanation in Causal Inference: Methods for Mediation and Interaction (Oxford University Press, 2015).",
]


def word_count(text: str) -> int:
    return len(re.findall(r"\b[\w-]+\b", text))


def add_labeled_image(doc: Document, label: str, image_path: Path,
                      max_width_cm: float = 16.5, max_height_cm: float = 22.0) -> None:
    """Keep a figure label and its image together during Word/PDF rendering."""
    with Image.open(image_path) as im:
        w, h = im.size
    width_cm = max_width_cm
    height_cm = width_cm * h / w
    if height_cm > max_height_cm:
        height_cm = max_height_cm
        width_cm = height_cm * w / h

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_after = Pt(8)
    label_run = p.add_run(label)
    base.font_run(label_run, 10.5, True, None, base.NAVY)
    label_run.add_break()
    image_run = p.add_run()
    image_run.add_picture(str(image_path), width=Cm(width_cm))


def apply_superscript_citations(doc: Document) -> None:
    """Render Nature-style numeric citations as superscript before the reference list."""
    citation_pattern = re.compile(r"(?<!\d\.)(?<=[\.\);,])(\d+(?:-\d+)?(?:,\d+(?:-\d+)?)*)")
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "References":
            break
        text = paragraph.text
        if not text or not citation_pattern.search(text):
            continue
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            continue
        if text.startswith(("Authors:", "Affiliations:", "Correspondence:", "Running title:", "Keywords:")):
            continue

        first_run = paragraph.runs[0] if paragraph.runs else None
        size = 10.5
        bold = None
        italic = None
        color = base.TEXT
        if first_run is not None:
            if first_run.font.size is not None:
                size = first_run.font.size.pt
            bold = first_run.font.bold
            italic = first_run.font.italic
            if first_run.font.color and first_run.font.color.rgb:
                color = first_run.font.color.rgb

        for run in list(paragraph.runs):
            run._element.getparent().remove(run._element)

        pos = 0
        for match in citation_pattern.finditer(text):
            if match.start() > pos:
                run = paragraph.add_run(text[pos:match.start()])
                base.font_run(run, size=size, bold=bold, italic=italic, color=color)
            ref_run = paragraph.add_run(match.group(1))
            base.font_run(ref_run, size=max(size - 2, 7.0), bold=bold, italic=italic, color=color)
            ref_run.font.superscript = True
            pos = match.end()
        if pos < len(text):
            run = paragraph.add_run(text[pos:])
            base.font_run(run, size=size, bold=bold, italic=italic, color=color)


def image_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    names = [
        "arialbd.ttf" if bold else "arial.ttf",
        "DejaVuSans-Bold.ttf" if bold else "DejaVuSans.ttf",
        "LiberationSans-Bold.ttf" if bold else "LiberationSans-Regular.ttf",
    ]
    for name in names:
        try:
            return ImageFont.truetype(name, size)
        except OSError:
            pass
    return ImageFont.load_default()


def compose_npj_two_panel(left_path: Path, right_path: Path, out_path: Path, left_title: str, right_title: str) -> None:
    left = Image.open(left_path).convert("RGB")
    right = Image.open(right_path).convert("RGB")
    resample = Image.Resampling.LANCZOS
    panel_w = 2300
    max_panel_h = 2700
    left.thumbnail((panel_w, max_panel_h), resample)
    right.thumbnail((panel_w, max_panel_h), resample)

    margin = 120
    gap = 120
    header_h = 170
    canvas_w = margin * 2 + panel_w * 2 + gap
    canvas_h = margin + header_h + max(left.height, right.height) + margin
    canvas = Image.new("RGB", (canvas_w, canvas_h), "white")
    draw = ImageDraw.Draw(canvas)
    label_font = image_font(70, True)
    title_font = image_font(42, True)
    text_color = (31, 41, 55)

    def add_panel(x: int, label: str, title: str, panel: Image.Image) -> None:
        draw.text((x, margin - 10), label, fill=text_color, font=label_font)
        draw.text((x + 115, margin + 6), title, fill=text_color, font=title_font)
        canvas.paste(panel, (x + (panel_w - panel.width) // 2, margin + header_h))

    add_panel(margin, "a)", left_title, left)
    add_panel(margin + panel_w + gap, "b)", right_title, right)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    canvas.save(out_path, dpi=(600, 600), quality=95)
    canvas.save(out_path.with_suffix(".tiff"), dpi=(600, 600), compression="tiff_lzw")


def normalize_image_for_upload(src: Path, png_dst: Path, tiff_dst: Path | None = None) -> None:
    image = Image.open(src).convert("RGB")
    png_dst.parent.mkdir(parents=True, exist_ok=True)
    image.save(png_dst, dpi=(600, 600), quality=95)
    if tiff_dst is not None:
        image.save(tiff_dst, dpi=(600, 600), compression="tiff_lzw")


def relabel_figure2_for_npj(src: Path, png_dst: Path, tiff_dst: Path | None = None) -> None:
    image = Image.open(src).convert("RGB")
    draw = ImageDraw.Draw(image)
    font = image_font(40, True)
    labels = ["a)", "b)", "c)", "d)", "e)", "f)"]
    boxes = [
        (106, 260, 176, 318),
        (1920, 260, 1990, 318),
        (106, 1634, 176, 1692),
        (1920, 1634, 1990, 1692),
        (106, 3008, 176, 3066),
        (1920, 3008, 1990, 3066),
    ]
    for label, (x1, y1, x2, y2) in zip(labels, boxes):
        draw.rounded_rectangle((x1 - 3, y1 - 2, x2 + 5, y2 + 2), radius=16, fill=(18, 31, 45))
        draw.text((x1 + 5, y1 + 6), label, fill="white", font=font)
    png_dst.parent.mkdir(parents=True, exist_ok=True)
    image.save(png_dst, dpi=(600, 600), quality=95)
    if tiff_dst is not None:
        image.save(tiff_dst, dpi=(600, 600), compression="tiff_lzw")


def prepare_support_files() -> None:
    NPJ_DST.mkdir(parents=True, exist_ok=True)
    base.copy_support_files()
    manifest = NPJ_DST / "SCI_English_submission_manifest.md"
    if manifest.exists():
        manifest.unlink()
    for redundant in (NPJ_DST / "supplementary_tables.xlsx", NPJ_DST / "source_data.xlsx"):
        if redundant.exists():
            redundant.unlink()

    upload = NPJ_DST / "figures_for_upload"
    upload.mkdir(exist_ok=True)
    compose_figure1_npj()
    relabel_figure2_for_npj(
        NPJ_DST / "main_figures" / "Figure_2_microbiome_diversity.png",
        upload / "Figure_2.png",
        upload / "Figure_2.tiff",
    )
    shutil.copy2(upload / "Figure_2.png", NPJ_DST / "main_figures" / "Figure_2_microbiome_diversity.png")
    compose_npj_two_panel(
        ROOT / "12_figures_tables_npj" / "yuting_style_visualizations" / "Y01_da_validation_balloonplot.png",
        ROOT / "12_figures_tables_npj" / "yuting_style_visualizations" / "Y06_key_taxa_dose_response_lollipop.png",
        upload / "Figure_3.png",
        "Taxon-level signal matrix",
        "Candidate taxa ranking",
    )
    shutil.copy2(upload / "Figure_3.png", NPJ_DST / "main_figures" / "Figure_3_recurrent_taxa_dose_response.png")
    compose_npj_two_panel(
        ROOT / "12_figures_tables_npj" / "yuting_style_visualizations" / "Y03_ecological_module_network.png",
        ROOT / "12_figures_tables_npj" / "yuting_style_visualizations" / "Y04_integrated_evidence_split_tile_heatmap.png",
        upload / "Figure_4.png",
        "Co-abundance modules",
        "Integrated evidence map",
    )
    shutil.copy2(upload / "Figure_4.png", NPJ_DST / "main_figures" / "Figure_4_ecological_modules_evidence.png")

    shutil.copy2(base.SRC / "supplementary_tables.xlsx", NPJ_DST / SUPP_DATA)
    shutil.copy2(base.SRC / "source_data.xlsx", NPJ_DST / SOURCE_DATA)


def prepare_npj_doc() -> Document:
    doc = base.prepare_doc()
    sec = doc.sections[0]
    sec.left_margin = Cm(2.1)
    sec.right_margin = Cm(2.1)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.35
    doc.styles["Normal"].font.size = Pt(11)
    return doc


def add_npj_table1(doc: Document) -> None:
    base.add_caption(
        doc,
        "Table 1. Selected weighted baseline characteristics across quartiles of the energy-adjusted dietary inflammatory index (E-DII).",
    )
    headers = ["Characteristic", "Q1", "Q2", "Q3", "Q4", "P"]
    rows = base.baseline_table_rows()
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    for j, h in enumerate(headers):
        base.fill_cell(table.cell(0, j), h)
    for i, row in enumerate(rows, 1):
        for j, val in enumerate(row):
            base.fill_cell(table.cell(i, j), val)
    base.format_table(table, font_size=7.2)
    base.add_para(
        doc,
        "Values are survey-weighted mean (SE) or weighted percentage (SE), unless otherwise indicated. "
        "The full baseline table by all eight dietary indices is provided in Supplementary Data 1.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        size=8.5,
        color=base.GRAY,
        first_indent=False,
        after=8,
        line_spacing=1.05,
    )


def add_title_page(doc: Document) -> None:
    base.add_title(doc, TITLE)
    base.add_para(
        doc,
        "Article type: Article",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=False,
        size=9.5,
        color=base.GRAY,
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.1
    author_parts = [
        ("Authors: ", False),
        ("Zhifeng Lin", False),
        ("1", True),
        (", Anping Liu", False),
        ("1", True),
        (", Jiajing Zhuang", False),
        ("2", True),
        (", Yiming Hu", False),
        ("3", True),
        (", Qingquan Chen", False),
        ("4", True),
        (", Zhijian Hu", False),
        ("1,5", True),
    ]
    for text, superscript in author_parts:
        run = p.add_run(text)
        base.font_run(run, 10.5)
        run.font.superscript = superscript
        if superscript:
            run.font.size = Pt(8)
    front_matter = [
        "Affiliations:",
        "1 Department of Epidemiology and Health Statistics, Fujian Provincial Key Laboratory of Environment Factors and Cancer, School of Public Health, Fujian Medical University, Fuzhou, China",
        "2 Fujian Medical University Union Hospital, Fuzhou, China",
        "3 National Center for Chronic and Noncommunicable Disease Control and Prevention, Chinese Center for Disease Control and Prevention, Beijing, China",
        "4 School of Biological Science and Medical Engineering and Beijing Advanced Innovation Centre for Biomedical Engineering, Beihang University, Beijing, China",
        "5 Key Laboratory of the Ministry of Education for Gastrointestinal Cancer, Fujian Medical University, Fuzhou, China",
        "Correspondence: Zhijian Hu (huzhijian@fjmu.edu.cn)",
        "Running title: Diet and oral microbiome ecology",
        "Keywords: oral microbiome; biofilm ecology; dietary inflammatory index; diet quality; NHANES; beta diversity",
    ]
    for item in front_matter:
        base.add_para(doc, item, first_indent=False, after=2, line_spacing=1.1)


def add_main_text(doc: Document) -> None:
    add_title_page(doc)

    base.add_heading(doc, "Abstract", 1)
    base.add_para(doc, ABSTRACT, first_indent=False)

    base.add_heading(doc, "Introduction", 1)
    intro = [
        "The oral microbiome is a biofilm-rich microbial ecosystem exposed directly to dietary substrates, saliva, dental surfaces, mucosal immune mediators and oral-hygiene behaviors.1-3 Diet can modify local nutrient availability, pH, oxidative stress, salivary composition and gingival inflammatory tone, all of which may plausibly influence oral microbial community structure. However, population-scale evidence linking overall diet quality and dietary inflammatory potential to oral microbiome ecology remains limited.",
        "Dietary indices provide a practical way to summarize complex diet patterns. The energy-adjusted Dietary Inflammatory Index was treated as the primary exposure because inflammatory potential is directly relevant to host-microbe ecology and energy adjustment reduces confounding by total intake.4 The unadjusted Dietary Inflammatory Index and six comparative indices were evaluated to distinguish inflammatory, guideline-based, plant-based and sustainability-oriented dietary dimensions.5-10",
        "Oral microbiome studies require ecological outcomes beyond alpha diversity. Within-sample richness or evenness may remain stable even when community composition, recurrent taxa or co-abundance modules shift subtly. We therefore evaluated alpha diversity, beta diversity, taxon-level centered log-ratio features, co-abundance modules and sensitivity analyses in a prespecified hierarchy. The goal was to determine whether small ecological signals recur across analytical layers, not to infer clinical causality.",
        "The NHANES 2009-2012 Oral Microbiome Project is an older data resource, but it remains one of the few nationally sampled public datasets linking oral 16S rRNA microbiome data with dietary recalls, oral-health phenotyping, medication information and complex survey design variables.11 We used this resource to examine whether dietary inflammatory potential and diet quality are associated with oral microbiome ecology, while explicitly treating non-causal oral-health association-path summaries as exploratory hypothesis generation.",
    ]
    for para in intro:
        base.add_para(doc, para)

    base.add_heading(doc, "Results", 1)
    results = [
        (
            "Analytic sample and diet-index architecture",
            "The oral microbiome metadata contained 9660 participants. Complete full-model analyses included 2799 participants with available dietary indices, survey design variables, covariates and microbiome outcomes; alpha-diversity models included 2656 participants. The dietary indices were correlated but not redundant. Principal-component analysis identified a dominant shared diet-quality axis explaining 66.3% of variance and a secondary axis explaining 12.1% that separated inflammatory and plant-based or sustainability-oriented dimensions (Fig. 1). Selected weighted baseline characteristics across energy-adjusted Dietary Inflammatory Index quartiles are shown in Table 1, and full baseline characteristics are provided in Supplementary Data 1."
        ),
        (
            "Alpha-diversity associations were weak",
            "Survey-weighted fully adjusted alpha-diversity models provided little evidence for large differences in within-sample richness or evenness. Associations across observed ASVs, Faith phylogenetic diversity, Shannon index and inverse Simpson index were small and did not survive false-discovery-rate correction. These results argue against a strong diet-related alpha-diversity phenotype in this cross-sectional sample."
        ),
        (
            "Beta-diversity showed small recurrent structural signals",
            "Beta-diversity analyses showed repeated but small community-structure signals across dietary indices and distance metrics. PERMANOVA R2 values ranged from 0.00031 to 0.00158; 17 of 32 tests met FDR < 0.10, but none met FDR < 0.05 (Fig. 2). PERMDISP was FDR-supported in several comparisons, indicating that some beta-diversity findings may include dispersion components rather than location shifts alone. Because distance-matrix analyses were not fully survey-weighted, these results are interpreted as structural ecological analyses rather than nationally representative effect estimates."
        ),
        (
            "Candidate taxa emerged through internal method-concordance",
            "Taxon-level centered log-ratio models identified 83 nominal diet-taxon associations and four exposure-domain FDR-supported associations, but no taxonomic feature met global FDR < 0.10. Internal method-concordance prioritized candidate taxa rather than validated biomarkers. Examples included Leptotrichia-related features in relation to the healthful plant-based diet index and Bifidobacterium in relation to the energy-adjusted Dietary Inflammatory Index (Fig. 3)."
        ),
        (
            "Co-abundance modules provided exploratory ecological context",
            "Co-abundance module analyses suggested higher-order ecological structure that was more coherent than many single-taxon findings. The M1 association with the healthful plant-based diet index was the strongest module signal (beta, 0.062; global FDR, 0.019). Larger modules, including M4 and M6, provided descriptive context for recurrent taxa and integrated evidence maps (Fig. 4). These modules are interpreted as exploratory co-abundance summaries, not as direct microbial interaction networks."
        ),
        (
            "Sensitivity and non-causal path summaries supported cautious hypothesis generation",
            "Sensitivity analyses generally preserved the direction of candidate taxa and module signals, including medication, dental-care and weighted-versus-unweighted comparisons. Oral-health path-coefficient summaries were retained in the Supplementary Information only as exploratory association-path analyses. Because diet, microbiome features and oral-health outcomes were measured cross-sectionally, these models cannot establish causal pathways, mediation or temporal order."
        ),
    ]
    for heading, body in results:
        base.add_heading(doc, heading, 2)
        base.add_para(doc, body)

    base.add_heading(doc, "Discussion", 1)
    discussion = [
        "This study found that dietary inflammatory potential and diet quality were associated with small but recurrent oral microbiome ecological signals in NHANES 2009-2012. The findings were not defined by large alpha-diversity differences or globally FDR-supported taxon-level discoveries. Instead, diet-related patterns appeared as modest beta-diversity structure, internally concordant candidate taxa and exploratory co-abundance modules.",
        "The magnitude of these associations was small. This is important for interpretation because oral microbiome structure is influenced by many stronger determinants, including dentition, periodontal status, smoking, age, medication use, oral hygiene and socioeconomic factors. The present results therefore should not be translated into clinical recommendations or microbial biomarkers. Their value is in defining candidate ecological patterns that can be tested in contemporary longitudinal cohorts, dietary intervention studies and functional microbiome analyses.",
        "The findings fit the scope of biofilm and microbiome ecology because they examine how a host-level environmental exposure, diet, relates to biofilm-associated oral microbial community structure. The lack of strong alpha-diversity findings is biologically plausible: dietary pressures may alter relative abundance, community dispersion or co-abundance structure without producing a large shift in richness.12 The recurrence of beta-diversity, candidate-taxon and module-level patterns supports a signal-convergence interpretation, but the evidence remains exploratory.",
        "Several limitations temper the conclusions. The analysis is cross-sectional, so temporal sequence cannot be established. NHANES 2009-2012 is older than contemporary datasets, although it remains a rare public national oral microbiome resource. Complete-case restriction reduced the analytic sample and may introduce selection bias. Supplementary Data 1 reports sample-construction counts, analysis-specific sample sizes and an included-versus-excluded descriptive comparison. Included participants were older and showed modest differences in poverty-income ratio, body mass index, race/ethnicity and smoking distribution compared with excluded participants. Dietary indices were derived from 24-hour dietary recalls, which may not fully capture usual intake. Distance-matrix analyses could not fully incorporate NHANES survey weights. Some PERMANOVA results may include dispersion effects. The public 16S data resolve taxa mainly at genus or family level and do not provide functional or strain-level inference.",
        "Overall, the results support the presence of small but consistent ecological associations between diet patterns and oral microbiome structure. They should be viewed as a foundation for replication and hypothesis generation rather than evidence for causal dietary modulation of oral biofilms or oral-health outcomes.",
    ]
    for para in discussion:
        base.add_para(doc, para)

    base.add_heading(doc, "Methods", 1)
    methods = [
        (
            "Study design and data source",
            "This cross-sectional secondary analysis used publicly available NHANES 2009-2010 and 2011-2012 data linked to the NHANES Oral Microbiome Project.11,13 Analyses were performed in May 2026. Participants were linked using SEQN across oral microbiome metadata, dietary recall files, demographic data, examination data, questionnaire files, oral-health variables, dental-care variables and prescription medication files."
        ),
        (
            "Ethics",
            "NHANES protocols were approved by the National Center for Health Statistics Research Ethics Review Board, and participants provided written informed consent. This secondary analysis used publicly available de-identified data and required no additional participant contact."
        ),
        (
            "Participants and missing data",
            "The source oral microbiome metadata contained 9660 participants. We used analysis-specific complete-case definitions because dietary indices, covariates, survey design variables and microbiome endpoints were not available for all participants. Complete full-model analyses included 2799 participants, and alpha-diversity analyses included 2656 participants. No imputation was applied to microbiome outcomes, dietary indices or model covariates. Analysis-specific denominators, sample-construction counts and an included-versus-excluded descriptive comparison are reported in Supplementary Data 1."
        ),
        (
            "Dietary exposures",
            "Dietary indices were constructed from NHANES 24-hour dietary recalls linked to the oral microbiome metadata. For participants with 2 recall days, nutrient and energy summaries used the 2-day average; if only 1 recall day was available, the available day was used. The 4-year dietary recall weight was constructed from the 2-day dietary weight divided by 2 for the two NHANES cycles. The primary exposure was the energy-adjusted Dietary Inflammatory Index.4 Secondary exposures were the unadjusted Dietary Inflammatory Index, Healthy Eating Index 2015, DASH score, alternate Mediterranean diet score, healthful plant-based diet index, unhealthful plant-based diet index and Planetary Health Diet Index.5-10 Continuous indices were standardized to a mean of 0 and standard deviation of 1 for regression models. Quartiles were used for descriptive analyses and distance-based community comparisons. Participants with missing exposure, covariate, design or endpoint variables were excluded from the corresponding complete-case model; no additional extreme-energy exclusion was applied beyond the source dietary-data availability rules. Component definitions and scoring directions are provided in Supplementary Data 1."
        ),
        (
            "Microbiome outcomes",
            "Public NHANES 16S rRNA oral microbiome outputs from oral specimens were used without reprocessing raw sequencing reads.11 The public annotation and distance-matrix files supplied taxonomy strings, alpha-diversity outputs and precomputed Bray-Curtis and UniFrac distances. Alpha-diversity outcomes included observed ASVs, Faith phylogenetic diversity, Shannon index and inverse Simpson index. Beta-diversity outcomes included Bray-Curtis, Jaccard, weighted UniFrac and unweighted UniFrac distances.14 The taxon matrix contained 1348 relative-abundance features before filtering. For centered log-ratio regression models, 102 taxonomic features with prevalence of at least 10% and mean relative abundance of at least 0.01% were retained. A pseudo-count of 2.375 x 10^-6 was added before log transformation, and each sample was centered by subtracting its row mean log abundance.18 The public taxonomy included features annotated to genus and higher unresolved taxonomic ranks; therefore these analyses are described as taxon-level."
        ),
        (
            "Covariates",
            "Models included demographic, socioeconomic, lifestyle, energy-intake, chronic disease, oral-health, dental-care and medication variables. These variables were selected to address major confounding pathways, while recognizing that some oral-health and dental-care variables may also lie on causal pathways or act as colliders. Therefore, oral-health extension and association-path summaries were interpreted cautiously and not used for causal inference."
        ),
        (
            "Survey design and weighting",
            "Descriptive analyses and alpha-diversity regression incorporated NHANES strata, primary sampling units and dietary sample weights.13,25 Two-cycle combined weights were divided by 2. PERMANOVA and PCoA analyses used precomputed distance matrices and were not interpreted as fully survey-weighted national estimates. This boundary was prespecified because standard complex-survey variance estimation is not directly compatible with the public precomputed distance matrices."
        ),
        (
            "Alpha-diversity models",
            "Survey-weighted linear regression estimated associations between standardized dietary indices and alpha-diversity outcomes. Restricted cubic spline models were used as nonlinear diagnostics.24 Estimates were reported as beta coefficients per 1-standard-deviation higher diet index, with uncertainty intervals and FDR values supplied in Supplementary Data 1."
        ),
        (
            "Beta-diversity and dispersion analyses",
            "PERMANOVA tested dietary-index quartile differences in distance-matrix structure.15 PERMDISP evaluated homogeneity of group dispersion,16 and constrained ordination summarized covariate-adjusted community structure using vegan.17 Results were interpreted using both PERMANOVA R2 and PERMDISP diagnostics."
        ),
        (
            "Taxon-level and module analyses",
            "Taxon-level models used centered log-ratio transformed abundance features.18 Internal method-concordance incorporated survey-weighted CLR regression, LEfSe ranking,19 and ANCOM-BC2 or MaAsLin-style sensitivity summaries.20-22 Co-abundance modules were inferred from Pearson correlations among Model 3 covariate-residualized CLR taxon abundances. Hierarchical clustering of 1 - absolute residual correlation was cut into 8 modules. Module scores were computed as first principal-component scores of the raw CLR abundances for taxa in each module and then related to dietary indices using NHANES survey-weighted linear models adjusted for Model 3 covariates. Network edges were displayed when absolute residual Pearson r was at least 0.20 and FDR was less than 0.05. Module networks are interpreted as co-abundance structures rather than direct ecological interactions."
        ),
        (
            "Multiple testing and evidence hierarchy",
            "Benjamini-Hochberg FDR correction was applied within prespecified analysis families, including alpha-diversity models, spline tests, beta-diversity metrics, taxon-level models, module associations, subgroup interactions, sensitivity analyses and exploratory non-causal oral-health association-path summaries.23 FDR < 0.05 was considered stronger statistical support, while FDR < 0.10 was treated as suggestive. Integrated evidence scores were descriptive synthesis tools and were not treated as inferential tests."
        ),
        (
            "Software",
            "Analyses were implemented using R and Python scripts organized by data preparation, exposure construction, covariate harmonization, alpha diversity, beta diversity, differential abundance, ecological modules, sensitivity analyses and figure generation. The analysis scripts will be made publicly available in a permanent repository before publication."
        ),
    ]
    for heading, body in methods:
        base.add_heading(doc, heading, 2)
        base.add_para(doc, body)

    base.add_heading(doc, "Data availability", 1)
    base.add_para(doc, "The raw NHANES data are publicly available from the National Center for Health Statistics, including the NHANES 2009-2010 and 2011-2012 public-use files and the Oral Microbiome Project outputs. The derived numeric tables supporting this manuscript are provided as Supplementary Data 1. Source data used for the figures are provided as the Source Data file.")

    base.add_heading(doc, "Code availability", 1)
    base.add_para(doc, "Custom R and Python scripts used to reproduce the derived tables, statistical analyses and figures will be made publicly available in a permanent repository before publication. The repository URL and archive DOI will be added to this statement when available.")

    base.add_heading(doc, "Acknowledgements", 1)
    base.add_para(doc, "This work was supported by the National Natural Science Foundation of China (No. 82473698), the Central government-led local science and technology development special project (Nos. 2019L3006 and 2020L3009), the Fujian Provincial Financial Special Project (23SCZZX004), and the Natural Science Foundation of Fujian Province (Nos. 2021J01726 and 2021J01733). The funders had no role in study design, data analysis, interpretation or writing of the manuscript.")

    base.add_heading(doc, "Author contributions", 1)
    base.add_para(doc, "Z.H. conceived the study and supervised the project. Z.L. and Y.H. conducted the formal analyses. Z.L. drafted the manuscript. A.L. and J.Z. revised the manuscript for important intellectual content. Y.H. and Q.C. curated the data. All authors interpreted the results, reviewed the manuscript and approved the final version.")

    base.add_heading(doc, "Competing interests", 1)
    base.add_para(doc, "The authors declare no competing interests.")

    base.add_heading(doc, "References", 1)
    for idx, ref in enumerate(NATURE_REFERENCES, 1):
        base.add_para(doc, f"{idx}. {ref}", first_indent=False, size=9.0, after=2, line_spacing=1.05)

    base.add_heading(doc, "Figure legends", 1)
    for label, text in FIGURE_LEGENDS:
        base.add_para(doc, f"{label}. {text}", first_indent=False, size=9.5, after=4, line_spacing=1.1)

    add_npj_table1(doc)

    for fig_name, src in [
        ("Figure 1", NPJ_DST / "figures_for_upload" / "Figure_1.png"),
        ("Figure 2", NPJ_DST / "figures_for_upload" / "Figure_2.png"),
        ("Figure 3", NPJ_DST / "figures_for_upload" / "Figure_3.png"),
        ("Figure 4", NPJ_DST / "figures_for_upload" / "Figure_4.png"),
    ]:
        add_labeled_image(doc, fig_name, src, max_width_cm=16.5, max_height_cm=22.0)


SUPP_FIGURE_FILES = [
    # Main-text panels are excluded from Supplementary Information to avoid duplicate display.
    ("Supplementary_Fig_S4_alpha_forest.png", "Extended forest plots of dietary indices and alpha-diversity outcomes."),
    ("Supplementary_Fig_S5_alpha_rcs.png", "Restricted cubic spline curves for dietary indices and alpha diversity."),
    ("Supplementary_Fig_S7_pcoa_quartiles.png", "Additional PCoA convex hulls by dietary-index quartile."),
    ("Supplementary_Fig_S9_da_validation.png", "Internal method-concordance summary for taxon-level findings."),
    ("Supplementary_Fig_S11_circular_taxa.png", "Circular visualization of recurrent marker taxa."),
    ("Supplementary_Fig_S12_upset_taxa.png", "UpSet plot of shared taxa across evidence domains."),
    ("Supplementary_Fig_S14_ecological_network.png", "Exploratory co-abundance module network."),
    ("Supplementary_Fig_S15_module_bridge.png", "Module-alpha diversity bridge summaries."),
    ("Supplementary_Fig_S16_community_type.png", "Dominant-taxon composition of community types."),
    ("Supplementary_Fig_S17_subgroup_interaction.png", "Subgroup and interaction analysis."),
    ("Supplementary_Fig_S18_med_dental_sensitivity.png", "Medication and dental-care sensitivity analysis."),
    ("Supplementary_Fig_S19_weighted_unweighted.png", "Weighted versus unweighted model comparison."),
    ("Supplementary_Fig_S21_path_coefficients.png", "Exploratory non-causal path-coefficient map."),
]

SUPP_FIGURE_LIST = [
    (f"Supplementary Fig. {i}", fname, caption)
    for i, (fname, caption) in enumerate(SUPP_FIGURE_FILES, start=1)
]


def build_main() -> None:
    doc = prepare_npj_doc()
    add_main_text(doc)
    apply_superscript_citations(doc)
    doc.save(NPJ_DST / MAIN_DOC)


def build_supplement() -> None:
    doc = prepare_npj_doc()
    base.add_title(
        doc,
        "Supplementary Information",
        "Dietary inflammatory potential and oral microbiome ecology in NHANES",
    )
    base.add_para(
        doc,
        "This file contains supplementary results, supplementary figure captions and embedded supplementary figures. "
        "All Methods descriptions are retained in the main manuscript to comply with npj Biofilms and Microbiomes Article format. "
        f"The full numeric tables are supplied separately as {SUPP_DATA}.",
        first_indent=False,
    )
    base.add_heading(doc, "Supplementary results", 1)
    notes = [
        "Supplementary Data 1 contains the complete numeric results for sample construction, included-versus-excluded participant comparison, exposure definitions, covariates, baseline characteristics, alpha-diversity models, beta-diversity analyses, taxon-level associations, recurrent taxa, ecological modules, subgroup analyses, sensitivity analyses, weighted-versus-unweighted comparisons and exploratory oral-health path-coefficient summaries.",
        "The supplementary figures are provided to document additional analytical context and robustness. Main-text figure panels are not repeated in the Supplementary Information. The supplementary figures should be interpreted alongside the main text hierarchy: principal beta-diversity analyses first, then secondary alpha-diversity and taxon-level analyses, followed by exploratory modules and non-causal oral-health association-path summaries.",
        "The association-path diagrams are retained only as hypothesis-generating visual summaries. They do not imply temporal sequence, causal mediation, causal pathways or clinical intervention effects.",
    ]
    for note in notes:
        base.add_para(doc, note)

    base.add_heading(doc, "Supplementary Data legend", 1)
    base.add_para(
        doc,
        f"Supplementary Data 1. Numeric tables supporting the manuscript. The workbook {SUPP_DATA} contains sheets corresponding to source data, included-versus-excluded participant comparison, baseline characteristics, alpha-diversity models, beta-diversity and PERMDISP analyses, taxon-level models, co-abundance modules, sensitivity analyses and exploratory non-causal oral-health path-coefficient summaries.",
        first_indent=False,
    )

    base.add_heading(doc, "Supplementary figure legends", 1)
    for label, fname, caption in SUPP_FIGURE_LIST:
        base.add_caption(doc, f"{label}. {caption}")
        doc.paragraphs[-1].paragraph_format.keep_with_next = True
        doc.paragraphs[-1].paragraph_format.keep_together = True
        base.insert_image(doc, NPJ_DST / "supplementary_figures" / fname, max_width_cm=16.7, max_height_cm=21.8)

    doc.save(NPJ_DST / SUPP_DOC)


def build_cover_letter() -> None:
    doc = prepare_npj_doc()
    base.add_title(doc, "Cover letter")
    paragraphs = [
        "Dear Editors,",
        "We are pleased to submit the Article entitled \"Dietary inflammatory potential and oral microbiome ecology in NHANES\" for consideration in npj Biofilms and Microbiomes.",
        "This manuscript uses the NHANES 2009-2012 Oral Microbiome Project to evaluate whether dietary inflammatory potential and multiple diet-quality indices are associated with oral microbiome ecological structure. The study fits the journal scope by examining host-microbiome ecology and biofilm-associated oral microbial communities in a population-based setting.",
        "The main finding is deliberately cautious: diet-related oral microbiome differences were small but consistent across beta-diversity structure, candidate taxa and exploratory co-abundance modules, while alpha-diversity findings were weak and causal interpretations were avoided. We believe this transparent framing will be useful for researchers designing contemporary longitudinal and functional oral microbiome studies.",
        "All authors have approved the manuscript. The authors declare no competing interests. Data availability and code availability statements are included in the manuscript.",
        "Sincerely,",
        "Zhijian Hu",
        "Corresponding author: huzhijian@fjmu.edu.cn",
    ]
    for para in paragraphs:
        base.add_para(doc, para, first_indent=False)
    doc.save(NPJ_DST / COVER_DOC)


def write_manifest_and_checklist() -> None:
    checklist = NPJ_DST / "NPJ_submission_checklist_and_notes.md"
    checklist.write_text(
        "# npj Biofilms and Microbiomes submission checklist\n\n"
        "Official requirements checked on 2026-05-24.\n\n"
        "## Package files\n"
        f"- Main manuscript: `{MAIN_DOC}`\n"
        f"- Supplementary Information: `{SUPP_DOC}` and rendered PDF in `pdf/`\n"
        f"- Supplementary Data: `{SUPP_DATA}`\n"
        f"- Source Data: `{SOURCE_DATA}`\n"
        f"- Cover letter: `{COVER_DOC}`\n"
        "- Figure upload files: `figures_for_upload/Figure_1-4.png` and TIFF copies\n\n"
        "## Journal-format checks applied\n"
        f"- Title word count: {word_count(TITLE)} words; no punctuation.\n"
        f"- Abstract word count: {word_count(ABSTRACT)} words; no subheadings.\n"
        "- Main structure: Title page, Abstract, Introduction, Results, Discussion, Methods, Data availability, Code availability, Acknowledgements, Author contributions, Competing interests, References, Figure legends.\n"
        "- Discussion has no separate Limitations or Conclusions subheading.\n"
        "- Methods are in the main manuscript, not in Supplementary Information.\n"
        "- Supplementary Information is combined into one separate file; PDF rendering is generated for submission; Supplementary Data 1 remains a separate Excel upload.\n"
        "- Main-text references have been inserted as sequential numeric citations and the reference list has been restyled to a Nature-compatible format.\n"
        "- Multi-panel Figures 1-4 use a), b), c) style labels where applicable.\n\n"
        "## Author actions still required\n"
        "- Confirm final author order, initials and affiliations before submission.\n"
        "- Confirm that all authors approve the author contributions, funding and competing-interests statements.\n"
        "- Consider depositing code in Zenodo, OSF or GitHub if the authors prefer a public repository link instead of reasonable-request access.\n"
        "- Complete the Nature Portfolio Reporting Summary if requested at revision or by the submission system.\n"
        "- Add a STROBE checklist for the cross-sectional observational design.\n"
        "- Confirm whether AI-assisted language or formatting support must be disclosed according to Nature Portfolio policy.\n",
        encoding="utf-8",
    )
    manifest = NPJ_DST / "NPJ_submission_manifest.md"
    manifest.write_text(
        "# npj Biofilms and Microbiomes submission package\n\n"
        "This package reorganizes the manuscript from the prior SCI/JNO drafts into the npj Biofilms and Microbiomes Article format. "
        "It emphasizes small but consistent ecological signals, moves causal-sounding path claims to exploratory Supplementary Information, "
        "and supplies figure files in manuscript-embedded and separate upload-ready forms.\n",
        encoding="utf-8",
    )


def main() -> None:
    prepare_support_files()
    build_main()
    build_supplement()
    build_cover_letter()
    write_manifest_and_checklist()
    print(NPJ_DST / MAIN_DOC)
    print(NPJ_DST / SUPP_DOC)
    print(NPJ_DST / COVER_DOC)
    print(NPJ_DST / SUPP_DATA)
    print(NPJ_DST / SOURCE_DATA)


if __name__ == "__main__":
    main()
